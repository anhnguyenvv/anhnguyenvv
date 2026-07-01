from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_harvard_cv(filepath):
    document = Document()
    
    # Margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    def add_heading(text, level=1):
        h = document.add_heading(text, level=level)
        h_font = h.style.font
        h_font.name = 'Times New Roman'
        h_font.bold = True
        h_font.color.rgb = None
        if level == 1:
            h_font.size = Pt(14)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            h_font.size = Pt(12)
        return h

    def add_bullet(text):
        p = document.add_paragraph(style='List Bullet')
        
        # We might have bold text in the bullet like "GraphRAG System:"
        # Simple parser for "Bold Text:"
        match = re.match(r"^([^:]+:)(.*)$", text)
        if match and len(match.group(1).split()) <= 4:
            run = p.add_run(match.group(1))
            run.bold = True
            p.add_run(match.group(2))
        else:
            p.add_run(text)
        return p

    def add_bold_run(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True
        return run

    # Header
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run("NGUYEN THI LAN ANH")
    name_run.bold = True
    name_run.size = Pt(16)
    
    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("Ho Chi Minh City, Vietnam | 0385 160 232 | anhnguyenvv0605@gmail.com | github.com/anhnguyenvv")
    
    # SUMMARY
    add_heading("SUMMARY", level=2)
    sum_p = document.add_paragraph("Computer Science graduate (HCMUS) specializing in end-to-end AI systems. Experienced in building GraphRAG pipelines and Multi-agent workflows at Newtech Shop. Expert in leveraging MCP to unify data from Odoo and BigQuery for automated analytics and internal QA. Passionate about LLM observability and agentic infrastructure.")
    
    # EDUCATION
    add_heading("EDUCATION", level=2)
    p = document.add_paragraph()
    add_bold_run(p, "VNUHCM – University of Science (HCMUS)")
    p.add_run(" | Ho Chi Minh City, Vietnam")
    p2 = document.add_paragraph("Bachelor of Computer Science | 2021 – 2025 | GPA: 3.38 / 4.0")
    
    # SKILLS
    add_heading("SKILLS & INTERESTS", level=2)
    skills = [
        ("Programming", "Python (proficient), SQL, Java"),
        ("AI & LLM", "Pytorch, RAG (LangChain), AI Agents (LangGraph), MCP, Prompt Engineering, LLM Monitoring, Ragas evaluation"),
        ("NLP & CV", "Transformer fine-tuning, embeddings, reranking, OCR, OpenCV, CUDA"),
        ("Data", "Qdrant, MySQL, Postgres, MongoDB, BigQuery, ETL pipelines"),
        ("Backend & DevOps", "FastAPI, Flask, REST APIs, Docker, Nginx, Git, Linux, Prometheus + Grafana"),
        ("Languages", "English – working proficiency; Japanese – JLPT N5"),
        ("Interests", "Multi-agent systems, applied LLM research, Vietnamese-language NLP")
    ]
    for category, details in skills:
        p = document.add_paragraph(style='List Bullet')
        add_bold_run(p, category + ": ")
        p.add_run(details)

    # EXPERIENCE
    add_heading("EXPERIENCE", level=2)
    
    exp1 = document.add_paragraph()
    add_bold_run(exp1, "Newtech Shop")
    exp1.add_run(" | Ho Chi Minh City, Vietnam")
    exp2 = document.add_paragraph()
    exp2.add_run("AI Agent Developer | Feb 2026 – Apr 2026")
    
    add_bullet("GraphRAG System: Developed a multi-agent GraphRAG chatbot for internal document QA, leveraging Neo4j and MCP to integrate structured data from Odoo ERP and BigQuery.")
    add_bullet("Marketing Automation: Engineered an automated Facebook Ads agent processing 10+ campaigns daily.")
    add_bullet("MCP Infrastructure: Managed an internal AI gateway routing 20+ LLMs (GPT-5, Claude) and operated 4+ MCP servers to unify data access across Odoo, BigQuery, and OmniChat.")
    add_bullet("Tech stack: Python, LangGraph, Neo4j, MCP, FastAPI, Next.js 14, MySQL 8.0, BigQuery, Pancake Pages API v2, Sapo, OmniChat, PowerBI, Anthropic Claude, Google Gemini, Docker Compose, Nginx.")
    
    # PROJECTS
    add_heading("PROJECTS", level=2)
    
    # Meeting Transcriber
    proj_mt = document.add_paragraph()
    add_bold_run(proj_mt, "Meeting Transcriber System")
    proj_mt.add_run(" | Personal Project | Mar 2026 – Present")
    add_bullet("Engineered an offline meeting transcription and diarization system using Faster-Whisper and Pyannote Audio.")
    add_bullet("Automated multilingual translation and meeting minutes generation via Qwen2.5 (Ollama).")
    
    # RAG Chatbot
    proj1 = document.add_paragraph()
    add_bold_run(proj1, "RAG Chatbot for FIT-HCMUS Academic Advising")
    proj1.add_run(" | github.com/anhnguyenvv/RAG_Chatbot")
    proj1_sub = document.add_paragraph("AI Engineer (Academic Project) | Jul 2025 – Dec 2025")
    
    add_bullet("Engineered hybrid RAG chatbot offering both Classic (retrieve → rerank → generate) and Agentic (ReAct) pipelines, achieving 92% context recall, 90% faithfulness, and 71% factual correctness on the Ragas benchmark.")
    add_bullet("Built end-to-end data pipeline covering FIT-website PDF crawling, LLM-based OCR, structure-aware chunking, embedding generation, and vector indexing.")
    add_bullet("Designed two-stage retrieval (BM25 + cross-encoder reranking) and a ReAct agent with custom search tools and persistent session memory; deployed full-stack with monitoring dashboards and 78 pytest test cases.")
    add_bullet("Tech stack: Python, LangChain, Qdrant, HuggingFace embeddings, Google Gemini, FastAPI, React + Vite, Tailwind, MongoDB, SQLite, Prometheus, Grafana.")
    
    # Voice Platform
    proj2 = document.add_paragraph()
    add_bold_run(proj2, "Vietnamese Text Processing & Labeling Platform")
    proj2.add_run(" | github.com/anhnguyenvv/vietnamese-text-analyzer")
    proj2_sub = document.add_paragraph("AI Engineer (Academic Project) | Jan 2025 – Jul 2025")
    
    add_bullet("Built web-based Vietnamese NLP platform integrating summarization, named entity recognition, POS tagging, sentiment analysis, classification, and text-to-speech; fine-tuned PhoBERT, ViSoBERT, and ViT5; trained Piper (VITS-based) TTS on a custom Vietnamese dataset.")
    add_bullet("Developed RESTful APIs and a modular backend integrated with a React frontend, including preprocessing pipelines for noisy real-world Vietnamese business data.")
    add_bullet("Tech stack: PyTorch, Hugging Face Transformers, PhoBERT, ViSoBERT, ViT5, Piper, Flask, React.js.")
    
    document.save(filepath)

if __name__ == "__main__":
    filepath = r"d:\project\protofilo\CV_NguyenThiLanAnh_AIEngineer_NLP_v2.docx"
    create_harvard_cv(filepath)
    print(f"CV saved to {filepath}")
